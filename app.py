import streamlit as st
from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled, NoTranscriptFound
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from dotenv import load_dotenv
from google import genai
from google.genai import types
from PIL import Image
import io, re, os, datetime, base64
import json
import requests
from docx import Document
from docx.shared import Pt, RGBColor
import markdown as md_lib

load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

# Get proxy from environment variables (set by Apache/system)
PROXY_URL = os.getenv("PROXY_URL") or os.getenv("http_proxy") or os.getenv("HTTP_PROXY")

st.set_page_config(
    page_title="VidBlog AI",
    page_icon="🎬",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# ── SESSION DEFAULTS ──
for k, v in [
    ("dark_mode", False),
    ("page", "generate"),          # "generate" | "history"
    ("blog_content", None),
    ("cover_bytes", None),
    ("blog_filename", None),
    ("img_filename", None),
    ("word_count", 0),
    ("blog_language", "English"),
    ("blog_tone", "Professional"),
    ("thumb_url", None),
    ("video_title", ""),
]:
    if k not in st.session_state:
        st.session_state[k] = v

# ── HANDLE NAV QUERY PARAMS EARLY (before theme tokens are set) ──
# (removed — nav now handled via session_state buttons directly)

dark = st.session_state.dark_mode

# ══════════════════════════════════════════
#  THEME TOKENS
# ══════════════════════════════════════════
if dark:
    BG        = "#07071a"
    CARD      = "#0f0f24"
    CARD2     = "#13132e"
    BORDER    = "rgba(139,92,246,0.18)"
    BORDER_S  = "rgba(255,255,255,0.07)"
    TEXT      = "#f1f5f9"
    TEXT2     = "#cbd5e1"
    MUTED     = "#64748b"
    INPUT_BG  = "#0a0a1e"
    GRAD      = "radial-gradient(ellipse 120% 55% at 50% 0%, rgba(124,58,237,0.18) 0%, transparent 65%)"
    SHADOW    = "0 0 0 1px rgba(255,255,255,0.05), 0 24px 64px rgba(0,0,0,0.65)"
    SHADOW_SM = "0 0 0 1px rgba(255,255,255,0.05), 0 4px 20px rgba(0,0,0,0.4)"
    CHIP      = "rgba(255,255,255,0.05)"
    CHIP_B    = "rgba(255,255,255,0.09)"
    SCROLL_TH = "rgba(139,92,246,0.45)"
    STEP_BG   = "rgba(255,255,255,0.04)"
    STEP_BR   = "rgba(255,255,255,0.09)"
    STEP_C    = "#334155"
    NAV_BG    = "rgba(7,7,26,0.85)"
    TAG_BG    = "rgba(124,58,237,0.1)"
    TAG_BR    = "rgba(124,58,237,0.22)"
    TAG_C     = "#c4b5fd"
    HIST_HOVER= "rgba(124,58,237,0.08)"
    EMPTY_C   = "#334155"
else:
    BG        = "#f0f2fa"
    CARD      = "#ffffff"
    CARD2     = "#f7f8fc"
    BORDER    = "rgba(124,58,237,0.2)"
    BORDER_S  = "rgba(0,0,0,0.07)"
    TEXT      = "#0f172a"
    TEXT2     = "#475569"
    MUTED     = "#94a3b8"
    INPUT_BG  = "#ffffff"
    GRAD      = "radial-gradient(ellipse 120% 55% at 50% 0%, rgba(124,58,237,0.07) 0%, transparent 65%)"
    SHADOW    = "0 0 0 1px rgba(0,0,0,0.06), 0 24px 64px rgba(0,0,0,0.09)"
    SHADOW_SM = "0 0 0 1px rgba(0,0,0,0.06), 0 4px 20px rgba(0,0,0,0.06)"
    CHIP      = "rgba(0,0,0,0.04)"
    CHIP_B    = "rgba(0,0,0,0.09)"
    SCROLL_TH = "rgba(124,58,237,0.35)"
    STEP_BG   = "rgba(0,0,0,0.04)"
    STEP_BR   = "rgba(0,0,0,0.1)"
    STEP_C    = "#94a3b8"
    NAV_BG    = "rgba(240,242,250,0.9)"
    TAG_BG    = "rgba(124,58,237,0.07)"
    TAG_BR    = "rgba(124,58,237,0.18)"
    TAG_C     = "#7c3aed"
    HIST_HOVER= "rgba(124,58,237,0.05)"
    EMPTY_C   = "#cbd5e1"


# ══════════════════════════════════════════
#  GLOBAL CSS
# ══════════════════════════════════════════
st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');
*,*::before,*::after{{box-sizing:border-box;margin:0;padding:0}}
html,body,[class*="css"]{{font-family:'Inter',sans-serif!important}}

::-webkit-scrollbar{{width:5px;height:5px}}
::-webkit-scrollbar-track{{background:transparent}}
::-webkit-scrollbar-thumb{{background:{SCROLL_TH};border-radius:99px}}

.stApp{{
  background:{BG}!important;
  background-image:{GRAD}!important;
  min-height:100vh;
}}
#MainMenu,footer,header{{visibility:hidden}}
.block-container{{
  padding:0 1.5rem 6rem!important;
  max-width:780px!important;
  margin:0 auto!important;
}}
[data-testid="collapsedControl"]{{display:none!important}}
[data-testid="stSidebar"]{{display:none!important}}

/* ── INPUTS ── */
.stTextInput>div>div>input{{
  background:{INPUT_BG}!important;
  border:1.5px solid {BORDER_S}!important;
  border-radius:12px!important;
  color:{TEXT}!important;
  padding:13px 16px!important;
  font-size:14px!important;
  font-family:'Inter',sans-serif!important;
  transition:border-color .2s,box-shadow .2s!important;
}}
.stTextInput>div>div>input:focus{{
  border-color:#7c3aed!important;
  box-shadow:0 0 0 3px rgba(124,58,237,.12)!important;
  outline:none!important;
}}
.stTextInput>div>div>input::placeholder{{color:{MUTED}!important}}
.stTextInput label{{display:none!important}}

/* ── SELECTBOX ── */
.stSelectbox label{{
  color:{MUTED}!important;font-size:11px!important;
  font-weight:600!important;letter-spacing:.8px!important;
  text-transform:uppercase!important;margin-bottom:6px!important;
}}
[data-baseweb="select"]>div{{
  background:{INPUT_BG}!important;
  border:1.5px solid {BORDER_S}!important;
  border-radius:12px!important;
  color:{TEXT}!important;font-size:14px!important;
  transition:border-color .2s!important;
}}
[data-baseweb="select"]>div:hover{{border-color:#7c3aed!important}}
[data-baseweb="popover"] ul{{
  background:{CARD}!important;
  border:1px solid {BORDER_S}!important;
  border-radius:12px!important;padding:6px!important;
}}
[role="option"]{{
  color:{TEXT2}!important;border-radius:8px!important;
  font-size:13px!important;padding:8px 12px!important;
}}
[role="option"]:hover{{background:rgba(124,58,237,.1)!important;color:{TEXT}!important}}
[aria-selected="true"]{{background:rgba(124,58,237,.12)!important;color:#c4b5fd!important}}

/* ── GENERATE BUTTON ── */
.gen-btn .stButton>button{{
  background:linear-gradient(135deg,#7c3aed,#4f46e5,#2563eb)!important;
  color:#fff!important;border:none!important;
  border-radius:12px!important;padding:14px 28px!important;
  font-size:15px!important;font-weight:700!important;
  width:100%!important;letter-spacing:.2px!important;
  box-shadow:0 8px 28px rgba(124,58,237,.45)!important;
  transition:all .25s ease!important;
}}
.gen-btn .stButton>button:hover{{
  transform:translateY(-2px)!important;
  box-shadow:0 14px 36px rgba(124,58,237,.55)!important;
}}

/* ── BUTTON TEXT COLOR FIX ── */
.stButton>button p {{
  color:{TEXT}!important;
  -webkit-text-fill-color:{TEXT}!important;
}}
.stButton>button:hover p,
.stButton>button:focus p,
.stButton>button:active p {{
  color:#ffffff!important;
  -webkit-text-fill-color:#ffffff!important;
}}
{"/* dark mode: override Streamlit white button bg */.stButton>button{background:"+CARD+"!important;border-color:"+BORDER_S+"!important;}.stButton>button:hover{background:rgba(124,58,237,.22)!important;border-color:#7c3aed!important;}" if dark else ""}

/* ── NAV ITEMS ── */
.nav-item:hover {{ color:{TEXT} !important; }}

/* ── DOWNLOAD BUTTONS ── */
.stDownloadButton>button{{
  background:{CHIP}!important;
  border:1.5px solid {CHIP_B}!important;
  color:{TEXT2}!important;border-radius:12px!important;
  font-weight:600!important;font-size:13px!important;
  width:100%!important;padding:11px 16px!important;
  transition:all .2s!important;
}}
.stDownloadButton>button:hover{{
  border-color:#7c3aed!important;
  background:rgba(124,58,237,.08)!important;
  color:#c4b5fd!important;
}}

/* ── ALERTS ── */
.stAlert{{border-radius:12px!important;font-size:13px!important}}
.stSpinner>div{{border-top-color:#7c3aed!important}}
.stImage img{{border-radius:14px!important}}
[data-testid="column"]{{padding:0 5px!important}}
[data-testid="column"]:first-child{{padding-left:0!important}}
[data-testid="column"]:last-child{{padding-right:0!important}}

/* ── TABS ── */
.stTabs [data-baseweb="tab-list"]{{
  background:{CARD2}!important;
  border-radius:12px!important;
  padding:4px!important;
  border:1px solid {BORDER_S}!important;
  gap:4px!important;
}}
.stTabs [data-baseweb="tab"]{{
  background:transparent!important;
  border-radius:9px!important;
  color:{TEXT2}!important;
  font-size:13px!important;font-weight:600!important;
  padding:8px 18px!important;
  border:none!important;
}}
.stTabs [aria-selected="true"]{{
  background:linear-gradient(135deg,#7c3aed,#4f46e5)!important;
  color:#fff!important;
  box-shadow:0 4px 12px rgba(124,58,237,.35)!important;
}}
.stTabs [data-baseweb="tab-highlight"]{{display:none!important}}
.stTabs [data-baseweb="tab-border"]{{display:none!important}}

/* ── RADIO BUTTONS ── */
div[data-testid="stRadio"] label p {{
  color: {TEXT} !important;
}}
div[data-testid="stRadio"] label span {{
  color: {TEXT} !important;
}}
</style>
""", unsafe_allow_html=True)

# ── COMPONENT CSS ──
st.markdown(f"""
<style>
.navbar{{
  display:flex;align-items:center;justify-content:space-between;
  padding:16px 0 12px;
}}
.nav-logo{{display:flex;align-items:center;gap:10px}}
.nav-icon{{
  width:34px;height:34px;
  background:linear-gradient(135deg,#7c3aed,#3b82f6);
  border-radius:10px;display:flex;align-items:center;
  justify-content:center;box-shadow:0 4px 14px rgba(124,58,237,.4);
  flex-shrink:0;
}}
.nav-title{{font-size:16px;font-weight:800;color:{TEXT};letter-spacing:-.3px}}
.nav-title span{{
  background:linear-gradient(90deg,#a78bfa,#60a5fa);
  -webkit-background-clip:text;-webkit-text-fill-color:transparent;
}}

.hero{{text-align:center;padding:36px 0 24px}}
.hero-badge{{
  display:inline-flex;align-items:center;gap:6px;
  background:{TAG_BG};border:1px solid {TAG_BR};
  color:{TAG_C};padding:5px 16px;border-radius:100px;
  font-size:10px;font-weight:700;letter-spacing:2.5px;
  text-transform:uppercase;margin-bottom:18px;
}}
.hero-title{{
  font-size:clamp(32px,6vw,50px);font-weight:900;
  color:{TEXT};line-height:1.05;letter-spacing:-2.5px;margin-bottom:14px;
}}
.g1{{background:linear-gradient(90deg,#a78bfa,#818cf8);-webkit-background-clip:text;-webkit-text-fill-color:transparent}}
.g2{{background:linear-gradient(90deg,#60a5fa,#34d399);-webkit-background-clip:text;-webkit-text-fill-color:transparent}}
.hero-sub{{color:{TEXT2};font-size:15px;line-height:1.7;max-width:460px;margin:0 auto 20px}}
.hero-tags{{display:flex;justify-content:center;gap:8px;flex-wrap:wrap}}
.hero-tag{{
  display:inline-flex;align-items:center;gap:5px;
  background:{CHIP};border:1px solid {CHIP_B};
  color:{TEXT2};padding:5px 13px;border-radius:100px;
  font-size:11px;font-weight:500;
}}

.steps-card{{
  background:{CARD};border:1px solid {BORDER_S};
  border-radius:16px;padding:16px 20px;
  margin-bottom:20px;box-shadow:{SHADOW_SM};
}}
.steps-inner{{display:flex;align-items:flex-start;justify-content:space-between;position:relative}}
.steps-line{{position:absolute;top:14px;left:14px;right:14px;height:1px;background:{BORDER_S};z-index:0}}
.step{{display:flex;flex-direction:column;align-items:center;gap:7px;flex:1;position:relative;z-index:1}}
.step-dot{{
  width:28px;height:28px;border-radius:50%;
  display:flex;align-items:center;justify-content:center;
  font-size:11px;font-weight:700;
  background:{STEP_BG};border:1.5px solid {STEP_BR};color:{STEP_C};
}}
.step.active .step-dot{{background:rgba(124,58,237,.2);border-color:#7c3aed;color:#c4b5fd;box-shadow:0 0 0 4px rgba(124,58,237,.12)}}
.step.done .step-dot{{background:rgba(52,211,153,.15);border-color:#34d399;color:#34d399}}
.step-lbl{{font-size:9px;font-weight:600;color:{STEP_C};text-align:center;white-space:nowrap;letter-spacing:.4px;text-transform:uppercase}}
.step.active .step-lbl{{color:#a78bfa}}
.step.done .step-lbl{{color:#6ee7b7}}

.fcard{{background:{CARD};border:1px solid {BORDER_S};border-radius:18px;padding:24px 26px 28px;box-shadow:{SHADOW}}}
.flabel{{display:flex;align-items:center;gap:8px;color:{MUTED};font-size:10.5px;font-weight:700;margin-bottom:8px;letter-spacing:1px;text-transform:uppercase}}
.fdivider{{height:1px;background:linear-gradient(90deg,transparent,{BORDER_S},transparent);margin:20px 0}}
.fsublabel{{color:{MUTED};font-size:10.5px;font-weight:700;letter-spacing:1px;text-transform:uppercase;margin-bottom:12px}}

.thumb-box{{
  border-radius:12px;overflow:hidden;
  border:1px solid {BORDER_S};margin-bottom:16px;
  position:relative;
}}
.thumb-box img{{width:100%;display:block}}
.thumb-overlay{{
  position:absolute;inset:0;
  background:linear-gradient(to top, rgba(0,0,0,0.7) 0%, transparent 50%);
  display:flex;align-items:flex-end;padding:12px;
}}
.thumb-title{{color:#fff;font-size:13px;font-weight:700;line-height:1.4}}

.stats{{display:flex;gap:8px;flex-wrap:wrap;margin:14px 0}}
.schip{{
  display:flex;align-items:center;gap:5px;
  background:{CHIP};border:1px solid {CHIP_B};
  border-radius:100px;padding:5px 12px;
  font-size:11px;color:{TEXT2};font-weight:500;
}}
.schip strong{{color:{TEXT};font-weight:700}}

.rcard{{background:{CARD};border:1px solid {BORDER_S};border-radius:18px;padding:26px 28px;margin-top:18px;box-shadow:{SHADOW}}}
.rhead{{display:flex;align-items:center;gap:14px;margin-bottom:20px;padding-bottom:16px;border-bottom:1px solid {BORDER_S}}}
.ricon{{width:42px;height:42px;flex-shrink:0;background:linear-gradient(135deg,#7c3aed,#3b82f6);border-radius:12px;display:flex;align-items:center;justify-content:center;font-size:19px;box-shadow:0 6px 18px rgba(124,58,237,.35)}}
.rtitle{{font-size:15px;font-weight:800;color:{TEXT}}}
.rsub{{font-size:12px;color:{MUTED};margin-top:3px}}

.blog{{color:{TEXT2};font-size:15px;line-height:1.85;padding:2px 0}}
.blog h1{{font-size:25px!important;font-weight:900!important;color:{TEXT}!important;margin:0 0 16px!important;letter-spacing:-.5px!important;line-height:1.2!important}}
.blog h2{{font-size:18px!important;font-weight:800!important;color:{TEXT}!important;margin:26px 0 10px!important;padding-bottom:8px!important;border-bottom:1px solid {BORDER_S}!important}}
.blog h3{{font-size:15px!important;font-weight:700!important;color:{TEXT}!important;margin:18px 0 8px!important}}
.blog p{{color:{TEXT2}!important;margin-bottom:13px!important}}
.blog ul,.blog ol{{padding-left:20px!important;margin-bottom:13px!important}}
.blog li{{color:{TEXT2}!important;margin-bottom:5px!important}}
.blog strong{{color:{TEXT}!important;font-weight:700!important}}
.blog blockquote{{border-left:3px solid #7c3aed!important;padding:10px 16px!important;margin:16px 0!important;background:rgba(124,58,237,.06)!important;border-radius:0 10px 10px 0!important;color:{MUTED}!important;font-style:italic!important}}
.blog hr{{border:none!important;border-top:1px solid {BORDER_S}!important;margin:22px 0!important}}
.blog code{{background:rgba(124,58,237,.1)!important;color:#c4b5fd!important;padding:2px 7px!important;border-radius:5px!important;font-size:13px!important}}

.cpbtn{{
  background:{CHIP};border:1.5px solid {CHIP_B};
  color:{TEXT2};padding:7px 15px;border-radius:9px;
  font-size:12px;font-weight:600;cursor:pointer;
  font-family:'Inter',sans-serif;transition:all .2s;
  float:right;margin-bottom:12px;
}}
.cpbtn:hover{{background:rgba(124,58,237,.1)!important;border-color:#7c3aed!important;color:#c4b5fd!important}}

.sbanner{{
  display:flex;align-items:center;gap:12px;
  background:rgba(52,211,153,.07);border:1px solid rgba(52,211,153,.18);
  border-radius:14px;padding:14px 18px;margin:16px 0 10px;
}}
.sbanner-icon{{font-size:22px;flex-shrink:0}}
.sbanner-title{{color:#6ee7b7;font-size:13px;font-weight:700}}
.sbanner-sub{{color:{MUTED};font-size:11px;margin-top:2px}}

/* HISTORY PAGE */
.page-header{{
  padding:28px 0 20px;
  border-bottom:1px solid {BORDER_S};
  margin-bottom:24px;
}}
.page-title{{font-size:22px;font-weight:900;color:{TEXT};letter-spacing:-.5px}}
.page-sub{{color:{TEXT2};font-size:13px;margin-top:4px}}

.hgrid{{display:grid;grid-template-columns:1fr 1fr;gap:14px;margin-top:4px}}
@media(max-width:600px){{.hgrid{{grid-template-columns:1fr}}}}

.hcard{{
  background:{CARD};border:1px solid {BORDER_S};
  border-radius:16px;overflow:hidden;
  transition:border-color .2s,transform .2s,box-shadow .2s;
  cursor:pointer;
}}
.hcard:hover{{
  border-color:rgba(124,58,237,.35);
  transform:translateY(-2px);
  box-shadow:0 12px 32px rgba(124,58,237,.15);
}}
.hcard-img{{width:100%;aspect-ratio:16/9;object-fit:cover;display:block}}
.hcard-img-placeholder{{
  width:100%;aspect-ratio:16/9;
  background:linear-gradient(135deg,rgba(124,58,237,.15),rgba(59,130,246,.1));
  display:flex;align-items:center;justify-content:center;
  font-size:32px;
}}
.hcard-body{{padding:14px 16px 16px}}
.hcard-title{{color:{TEXT};font-size:13px;font-weight:700;line-height:1.4;margin-bottom:5px}}
.hcard-meta{{display:flex;gap:8px;flex-wrap:wrap;margin-bottom:10px}}
.hcard-chip{{
  background:{CHIP};border:1px solid {CHIP_B};
  color:{MUTED};font-size:10px;font-weight:600;
  padding:3px 9px;border-radius:100px;
}}
.hcard-preview{{color:{TEXT2};font-size:11.5px;line-height:1.6;max-height:52px;overflow:hidden}}

.empty-state{{
  text-align:center;padding:72px 24px;
}}
.empty-icon{{font-size:52px;margin-bottom:16px}}
.empty-title{{color:{TEXT};font-size:17px;font-weight:700;margin-bottom:8px}}
.empty-sub{{color:{MUTED};font-size:13px;line-height:1.6}}

.del-btn .stButton>button{{
  background:rgba(239,68,68,.08)!important;
  border:1.5px solid rgba(239,68,68,.2)!important;
  color:#f87171!important;border-radius:10px!important;
  font-size:12px!important;font-weight:600!important;
  padding:7px 14px!important;width:100%!important;
  transition:all .2s!important;
}}
.del-btn .stButton>button:hover{{
  background:rgba(239,68,68,.15)!important;
  border-color:rgba(239,68,68,.4)!important;
}}

.view-btn .stButton>button{{
  background:rgba(124,58,237,.1)!important;
  border:1.5px solid rgba(124,58,237,.25)!important;
  color:#c4b5fd!important;border-radius:10px!important;
  font-size:12px!important;font-weight:600!important;
  padding:7px 14px!important;width:100%!important;
  transition:all .2s!important;
}}
.view-btn .stButton>button:hover{{
  background:rgba(124,58,237,.18)!important;
  border-color:rgba(124,58,237,.45)!important;
}}

.cta-footer{{
  text-align:center;margin-top:22px;padding:20px;
  background:{CARD};border:1px solid {BORDER_S};border-radius:16px;
}}
.cta-footer p{{color:{TEXT2};font-size:13px;margin-bottom:4px}}
.cta-footer span{{color:{MUTED};font-size:12px}}
</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════
SAVE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "outputs")
os.makedirs(SAVE_DIR, exist_ok=True)

def md_to_html(md_text: str, title: str = "Blog") -> bytes:
    body = md_lib.markdown(md_text, extensions=["extra", "nl2br"])
    html = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<title>{title}</title>
<style>body{{font-family:Georgia,serif;max-width:800px;margin:40px auto;padding:0 20px;line-height:1.8;color:#1a1a1a}}
h1{{font-size:2em;margin-bottom:.5em}}h2{{font-size:1.4em;margin-top:1.5em}}
p{{margin-bottom:1em}}blockquote{{border-left:4px solid #7c3aed;padding:8px 16px;color:#555;margin:16px 0}}
code{{background:#f3f0ff;padding:2px 6px;border-radius:4px}}</style>
</head><body>{body}</body></html>"""
    return html.encode("utf-8")

def md_to_docx(md_text: str, title: str = "Blog") -> bytes:
    doc = Document()
    doc.core_properties.title = title
    for line in md_text.split("\n"):
        line = line.strip()
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def extract_video_id(url: str):
    for p in [r"(?:v=)([a-zA-Z0-9_-]{11})", r"(?:youtu\.be/)([a-zA-Z0-9_-]{11})",
              r"(?:embed/)([a-zA-Z0-9_-]{11})", r"(?:shorts/)([a-zA-Z0-9_-]{11})"]:
        m = re.search(p, url)
        if m: return m.group(1)
    return None

def get_thumbnail_url(vid: str) -> str:
    return f"https://img.youtube.com/vi/{vid}/maxresdefault.jpg"

def fetch_transcript_with_proxy(vid: str):
    """Fetch transcript using system proxy (Apache/environment)."""
    import time
    from urllib.request import ProxyHandler, build_opener
    
    # Build proxy dict from environment or explicit config
    proxies = {}
    if PROXY_URL:
        proxies = {
            "http": PROXY_URL,
            "https": PROXY_URL,
        }
    
    ytt = YouTubeTranscriptApi()
    
    # Configure requests session with proxy
    if proxies:
        session = requests.Session()
        session.proxies.update(proxies)
        # Inject session into transcript API
        ytt.http_client = session
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            tlist = ytt.list(vid)
            try:
                transcript = tlist.find_transcript(['en','en-US','en-GB','hi','ur'])
            except NoTranscriptFound:
                transcript = next(iter(tlist))
            return transcript.fetch()
        except Exception as e:
            error_msg = str(e).lower()
            # Check if it's an IP blocking error
            is_ip_blocked = any(x in error_msg for x in ['ip', 'blocked', 'request', 'forbidden'])
            
            if attempt < max_retries - 1 and is_ip_blocked:
                wait_time = 2 ** attempt  # exponential backoff: 1s, 2s, 4s
                st.info(f"⏳ Retrying in {wait_time}s... (attempt {attempt+1}/{max_retries})")
                time.sleep(wait_time)
            else:
                raise

def render_steps(active=None, done=None):
    done = done or []
    items = [("1","Transcript"),("2","Context"),("3","Blog"),("4","Cover"),("5","Done")]
    dots = ""
    for num, lbl in items:
        if lbl in done:    cls, sym = "step done",   "✓"
        elif lbl == active: cls, sym = "step active", num
        else:               cls, sym = "step",         num
        dots += f'<div class="{cls}"><div class="step-dot">{sym}</div><div class="step-lbl">{lbl}</div></div>'
    return f'<div class="steps-card"><div class="steps-inner"><div class="steps-line"></div>{dots}</div></div>'

@st.cache_data(ttl=60, show_spinner=False)
def load_history():
    """Return list of dicts with blog metadata, newest first. Cached for 60s."""
    items = []
    for fn in sorted(os.listdir(SAVE_DIR), reverse=True):
        if not fn.endswith(".md"):
            continue
        path = os.path.join(SAVE_DIR, fn)
        try:
            with open(path, "r", encoding="utf-8") as f:
                content = f.read()
            title = next((l.lstrip("# ").strip() for l in content.splitlines() if l.startswith("#")), fn)
            ts = fn.replace("blog_", "").replace(".md", "")
            try:
                ts_label = datetime.datetime.strptime(ts, "%Y%m%d_%H%M%S").strftime("%d %b %Y · %I:%M %p")
            except:
                ts_label = ts
            img_fn = fn.replace("blog_", "cover_").replace(".md", ".png")
            img_path = os.path.join(SAVE_DIR, img_fn)
            meta_fn = fn.replace("blog_", "meta_").replace(".md", ".json")
            meta_path = os.path.join(SAVE_DIR, meta_fn)
            meta = {}
            if os.path.exists(meta_path):
                with open(meta_path, "r", encoding="utf-8") as mf:
                    meta = json.load(mf)
            word_count = len(content.split())
            items.append({
                "fn": fn,
                "path": path,
                "title": title,
                "ts_label": ts_label,
                "ts_raw": ts,
                "content": content,
                "img_path": img_path if os.path.exists(img_path) else None,
                "img_fn": img_fn,
                "word_count": word_count,
                "language": meta.get("language", "English"),
                "tone": meta.get("tone", "Professional"),
                "hashtags": meta.get("hashtags", []),
                "read_time": max(1, word_count // 200),
            })
        except:
            pass
    return items


@st.cache_data(ttl=300, show_spinner=False)
def get_thumb_b64(img_path: str) -> str:
    """Resize image to thumbnail and return base64 — cached per path."""
    try:
        img = Image.open(img_path)
        img.thumbnail((480, 270))
        buf = io.BytesIO()
        img.save(buf, format="WEBP", quality=70)
        return base64.b64encode(buf.getvalue()).decode()
    except:
        return ""

def save_meta(ts: str, language: str, tone: str, hashtags: list = None):
    meta_fn = f"meta_{ts}.json"
    meta_data = {"language": language, "tone": tone}
    if hashtags:
        meta_data["hashtags"] = hashtags
    with open(os.path.join(SAVE_DIR, meta_fn), "w", encoding="utf-8") as f:
        json.dump(meta_data, f)

def generate_hashtags(blog_content: str, blog_title: str, llm) -> list:
    """Generate relevant hashtags for the blog post using Gemini."""
    try:
        prompt = f"""Based on this blog post, generate 10-15 relevant, trending hashtags.

Blog Title: {blog_title}

Blog Content (first 1000 chars):
{blog_content[:1000]}

Requirements:
- Return ONLY hashtags, one per line
- Start each with #
- No explanations or extra text
- Mix popular and niche hashtags
- Make them relevant to the content
- Keep them concise (2-3 words max per hashtag)"""
        
        response = llm.invoke(prompt)
        hashtags = [tag.strip() for tag in response.content.split('\n') if tag.strip().startswith('#')]
        return hashtags[:15]  # Return max 15 hashtags
    except Exception as e:
        st.warning(f"⚠️ Hashtag generation skipped: {e}")
        return []

TONE_MAP = {
    "Professional":      "formal, authoritative, and polished",
    "Casual & Friendly": "conversational, warm, and approachable",
    "Educational":       "clear, informative, and easy to understand",
    "Storytelling":      "narrative-driven, engaging, and vivid",
    "Technical":         "precise, detailed, and technically accurate",
    "Persuasive":        "compelling, motivating, and action-oriented",
}


# ══════════════════════════════════════════
#  NAVBAR  (always visible)
# ══════════════════════════════════════════
if not GEMINI_API_KEY:
    st.error("⚠️ GEMINI_API_KEY not found in .env file.")
    st.stop()

hist_count = len([f for f in os.listdir(SAVE_DIR) if f.endswith(".md")])
is_gen  = st.session_state.page == "generate"
is_hist = st.session_state.page == "history"
_theme_icon  = "🌙" if not dark else "☀️"
_theme_label = "Dark" if not dark else "Light"

# ── NAVBAR CSS ──
_gen_active_css  = f"background:rgba(124,58,237,.1)!important;border-color:#7c3aed!important;color:#7c3aed!important;" if is_gen  else ""
_hist_active_css = f"background:rgba(124,58,237,.1)!important;border-color:#7c3aed!important;color:#7c3aed!important;" if is_hist else ""

st.markdown(f"""
<style>
/* ── NAV BUTTONS BASE ── */
.nav-col .stButton>button,
.nav-btn-active .stButton>button {{
  height:38px!important;
  padding:0 18px!important;
  font-size:13px!important;
  font-weight:600!important;
  border-radius:10px!important;
  border:1.5px solid {BORDER_S}!important;
  background:{CARD}!important;
  color:{TEXT}!important;
  box-shadow:none!important;
  font-family:'Inter',sans-serif!important;
  transition:border-color .18s,background .18s,color .18s,box-shadow .18s!important;
  white-space:nowrap!important;
  width:100%!important;
  line-height:1!important;
}}

/* ── NAV BUTTON TEXT (Streamlit wraps in <p>) ── */
.nav-col .stButton>button p,
.nav-btn-active .stButton>button p {{
  color:{TEXT}!important;
  -webkit-text-fill-color:{TEXT}!important;
  font-size:13px!important;
  font-weight:600!important;
  margin:0!important;
}}

/* ── HOVER / FOCUS / ACTIVE ── */
.nav-col .stButton>button:hover,
.nav-col .stButton>button:focus,
.nav-col .stButton>button:active {{
  border-color:#7c3aed!important;
  background:rgba(124,58,237,.22)!important;
  box-shadow:0 0 0 3px rgba(124,58,237,.15)!important;
  color:#ffffff!important;
  -webkit-text-fill-color:#ffffff!important;
}}
.nav-col .stButton>button:hover p,
.nav-col .stButton>button:focus p,
.nav-col .stButton>button:active p {{
  color:#ffffff!important;
  -webkit-text-fill-color:#ffffff!important;
}}

/* ── ACTIVE PAGE BUTTON ── */
.nav-btn-active .stButton>button {{
  background:rgba(124,58,237,.18)!important;
  border-color:#7c3aed!important;
  color:#ffffff!important;
  -webkit-text-fill-color:#ffffff!important;
}}
.nav-btn-active .stButton>button p {{
  color:#ffffff!important;
  -webkit-text-fill-color:#ffffff!important;
}}
.nav-btn-active .stButton>button:hover,
.nav-btn-active .stButton>button:focus,
.nav-btn-active .stButton>button:active {{
  background:rgba(124,58,237,.30)!important;
  border-color:#a78bfa!important;
  color:#ffffff!important;
  -webkit-text-fill-color:#ffffff!important;
}}
.nav-btn-active .stButton>button:hover p,
.nav-btn-active .stButton>button:focus p,
.nav-btn-active .stButton>button:active p {{
  color:#ffffff!important;
  -webkit-text-fill-color:#ffffff!important;
}}
</style>
""", unsafe_allow_html=True)

# ── NAVBAR ── logo + buttons in one row
_logo_col, _c2, _c3, _c4 = st.columns([2.5, 1, 1, 1])

with _logo_col:
    _logo_path = os.path.join(os.path.dirname(__file__), "Prabisha_logo.png")
    if os.path.exists(_logo_path):
        with open(_logo_path, "rb") as _lf:
            _logo_b64 = base64.b64encode(_lf.read()).decode()
        _logo_html = '<div style="display:flex;align-items:center;gap:12px;height:38px"><img src="data:image/png;base64,' + _logo_b64 + '" style="height:52px;object-fit:contain;" alt="Prabisha Consulting"><span style="font-size:16px;font-weight:800;background:linear-gradient(90deg,#7c3aed,#3b82f6);-webkit-background-clip:text;-webkit-text-fill-color:transparent;letter-spacing:-0.5px;white-space:nowrap">Video to Blog</span></div>'
        st.markdown(_logo_html, unsafe_allow_html=True)
    else:
        st.markdown(f"""
        <div style="display:flex;align-items:center;height:38px">
          <div class="nav-logo">
            <div class="nav-icon">
              <svg width="14" height="14" viewBox="0 0 20 20" fill="none">
                <polygon points="4,3 4,17 16,10" fill="white"/>
              </svg>
            </div>
            <span class="nav-title">VidBlog <span>AI</span></span>
          </div>
        </div>
        """, unsafe_allow_html=True)

# push logo to left by using spacer; buttons go right
with _c2:
    _gen_bg = "rgba(124,58,237,.22)" if is_gen else CARD
    _gen_bc = "#7c3aed" if is_gen else BORDER_S
    _gen_tc = "#7c3aed" if is_gen else TEXT
    st.markdown(f"""
    <style>
    #nav_gen_wrap button {{ background:{_gen_bg}!important; border:1.5px solid {_gen_bc}!important; height:38px!important; border-radius:10px!important; font-size:13px!important; font-weight:600!important; width:100%!important; }}
    #nav_gen_wrap button p {{ color:{_gen_tc}!important; -webkit-text-fill-color:{_gen_tc}!important; }}
    #nav_gen_wrap button:hover {{ background:rgba(124,58,237,.35)!important; border-color:#a78bfa!important; }}
    #nav_gen_wrap button:hover p {{ color:#ffffff!important; -webkit-text-fill-color:#ffffff!important; }}
    </style>
    <div id="nav_gen_wrap">
    """, unsafe_allow_html=True)
    if st.button("✨ Generate", key="nav_gen", width='stretch'):
        st.session_state.page = "generate"
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

with _c3:
    _hist_bg = "rgba(124,58,237,.22)" if is_hist else CARD
    _hist_bc = "#7c3aed" if is_hist else BORDER_S
    _hist_tc = "#7c3aed" if is_hist else TEXT
    st.markdown(f"""
    <style>
    #nav_hist_wrap button {{ background:{_hist_bg}!important; border:1.5px solid {_hist_bc}!important; height:38px!important; border-radius:10px!important; font-size:13px!important; font-weight:600!important; width:100%!important; }}
    #nav_hist_wrap button p {{ color:{_hist_tc}!important; -webkit-text-fill-color:{_hist_tc}!important; }}
    #nav_hist_wrap button:hover {{ background:rgba(124,58,237,.35)!important; border-color:#a78bfa!important; }}
    #nav_hist_wrap button:hover p {{ color:#ffffff!important; -webkit-text-fill-color:#ffffff!important; }}
    </style>
    <div id="nav_hist_wrap">
    """, unsafe_allow_html=True)
    if st.button(f"📂 History", key="nav_hist", width='stretch'):
        st.session_state.page = "history"
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

with _c4:
    st.markdown(f"""
    <style>
    #nav_theme_wrap button {{ background:{CARD}!important; border:1.5px solid {BORDER_S}!important; height:38px!important; border-radius:10px!important; font-size:13px!important; font-weight:600!important; width:100%!important; }}
    #nav_theme_wrap button:hover {{ background:rgba(124,58,237,.22)!important; border-color:#7c3aed!important; }}
    </style>
    <div id="nav_theme_wrap">
    """, unsafe_allow_html=True)
    if st.button(f"{_theme_icon} {_theme_label}", key="nav_theme", width='stretch'):
        st.session_state.dark_mode = not st.session_state.dark_mode
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

st.markdown(f'<div style="height:1px;background:{BORDER_S};margin:6px 0 0"></div>', unsafe_allow_html=True)


# ══════════════════════════════════════════
#  PAGE: HISTORY
# ══════════════════════════════════════════
if st.session_state.page == "history":
    history = load_history()

    st.markdown(f"""
    <div class="page-header">
      <div class="page-title">📂 Blog History</div>
      <div class="page-sub">{len(history)} blog{"s" if len(history) != 1 else ""} saved locally in <code>outputs/</code></div>
    </div>
    """, unsafe_allow_html=True)

    if not history:
        st.markdown(f"""
        <div class="empty-state">
          <div class="empty-icon">📭</div>
          <div class="empty-title">No blogs yet</div>
          <div class="empty-sub">Generate your first blog from the Generate tab.<br>All blogs are saved automatically.</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        # ── VIEW SINGLE BLOG ──
        if "view_blog" in st.session_state and st.session_state.view_blog:
            vb = st.session_state.view_blog
            # Back button
            st.markdown('<div class="navbtn">', unsafe_allow_html=True)
            if st.button("← Back to History", key="back_hist"):
                st.session_state.view_blog = None
                st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)
            st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)

            # Cover image
            if vb["img_path"] and os.path.exists(vb["img_path"]):
                st.image(vb["img_path"], width='stretch')
                st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)

            # Stats
            st.markdown(f"""
            <div class="stats">
              <div class="schip">✍️ <strong>{vb["word_count"]:,} words</strong></div>
              <div class="schip">⏱️ <strong>~{vb["read_time"]} min read</strong></div>
              <div class="schip">🌐 <strong>{vb["language"]}</strong></div>
              <div class="schip">🎨 <strong>{vb["tone"]}</strong></div>
              <div class="schip">🕐 <strong>{vb["ts_label"]}</strong></div>
            </div>
            """, unsafe_allow_html=True)

            # Hashtags
            if vb.get("hashtags"):
                hashtags_html = '<div style="margin-top:14px;padding:12px 14px;background:' + CARD2 + ';border:1px solid ' + BORDER_S + ';border-radius:12px">'
                hashtags_html += '<div style="color:' + MUTED + ';font-size:10px;font-weight:700;letter-spacing:1px;text-transform:uppercase;margin-bottom:8px">🏷️ Hashtags</div>'
                hashtags_html += '<div style="display:flex;flex-wrap:wrap;gap:6px">'
                for tag in vb["hashtags"]:
                    hashtags_html += '<span style="background:' + CHIP + ';border:1px solid ' + CHIP_B + ';color:' + TEXT2 + ';padding:4px 10px;border-radius:100px;font-size:11px;font-weight:500;cursor:pointer" onclick="navigator.clipboard.writeText(\'' + tag + '\')">' + tag + '</span>'
                hashtags_html += '</div></div>'
                st.markdown(hashtags_html, unsafe_allow_html=True)

            # Blog content in tabs
            tab_preview, tab_raw = st.tabs(["📖 Preview", "📝 Raw Markdown"])
            with tab_preview:
                st.markdown('<div class="blog">', unsafe_allow_html=True)
                st.markdown(vb["content"])
                st.markdown('</div>', unsafe_allow_html=True)
            with tab_raw:
                st.code(vb["content"], language="markdown")

            # Downloads
            st.markdown(f"""
            <div style="height:1px;background:{BORDER_S};margin:20px 0 14px"></div>
            <div style="color:{MUTED};font-size:10.5px;font-weight:700;letter-spacing:1px;text-transform:uppercase;margin-bottom:12px">⬇️ Download</div>
            """, unsafe_allow_html=True)

            _hbase = vb["fn"].replace(".md", "")
            _has_img = vb["img_path"] and os.path.exists(vb["img_path"])
            _cols = st.columns(4) if _has_img else st.columns(3)
            with _cols[0]:
                st.download_button("📄 .md", data=vb["content"],
                    file_name=vb["fn"], mime="text/markdown",
                    width='stretch', key="hist_dl_blog")
            with _cols[1]:
                st.download_button("🌐 .html", data=md_to_html(vb["content"], _hbase),
                    file_name=_hbase+".html", mime="text/html",
                    width='stretch', key="hist_dl_html")
            with _cols[2]:
                st.download_button("📝 .docx", data=md_to_docx(vb["content"], _hbase),
                    file_name=_hbase+".docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    width='stretch', key="hist_dl_docx")
            if _has_img:
                with _cols[3]:
                    with open(vb["img_path"], "rb") as imgf:
                        img_bytes = imgf.read()
                    st.download_button("🖼️ .png", data=img_bytes,
                        file_name=vb["img_fn"], mime="image/png",
                        width='stretch', key="hist_dl_img")

        else:
            # ── GRID VIEW ──
            HIST_PAGE_SIZE = 10
            if "hist_page" not in st.session_state:
                st.session_state.hist_page = 0

            total = len(history)
            total_pages = max(1, (total + HIST_PAGE_SIZE - 1) // HIST_PAGE_SIZE)
            # clamp page index after deletions
            if st.session_state.hist_page >= total_pages:
                st.session_state.hist_page = total_pages - 1

            page_items = history[
                st.session_state.hist_page * HIST_PAGE_SIZE :
                (st.session_state.hist_page + 1) * HIST_PAGE_SIZE
            ]

            # Render cards in 2-column grid using HTML
            cards_html = '<div class="hgrid">'
            for item in page_items:
                if item["img_path"]:
                    b64 = get_thumb_b64(item["img_path"])
                    img_tag = f'<img class="hcard-img" src="data:image/webp;base64,{b64}" alt="cover">' if b64 else '<div class="hcard-img-placeholder">📝</div>'
                else:
                    img_tag = '<div class="hcard-img-placeholder">📝</div>'

                short_title = (item["title"][:55] + "…") if len(item["title"]) > 55 else item["title"]
                preview = item["content"][:160].replace('<','&lt;').replace('>','&gt;').replace('\n',' ')

                cards_html += f"""
                <div class="hcard">
                  {img_tag}
                  <div class="hcard-body">
                    <div class="hcard-title">{short_title}</div>
                    <div class="hcard-meta">
                      <span class="hcard-chip">🕐 {item["ts_label"]}</span>
                      <span class="hcard-chip">🌐 {item["language"]}</span>
                      <span class="hcard-chip">⏱️ {item["read_time"]}m read</span>
                    </div>
                    <div class="hcard-preview">{preview}…</div>
                  </div>
                </div>"""
            cards_html += '</div>'
            st.markdown(cards_html, unsafe_allow_html=True)

            # Pagination controls
            if total_pages > 1:
                st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)
                pg_cols = st.columns([1, 2, 1])
                with pg_cols[0]:
                    if st.button("← Prev", disabled=st.session_state.hist_page == 0, width='stretch', key="hist_prev"):
                        st.session_state.hist_page -= 1
                        st.rerun()
                with pg_cols[1]:
                    st.markdown(f"<div style='text-align:center;color:{MUTED};font-size:12px;padding-top:8px'>Page {st.session_state.hist_page+1} of {total_pages}</div>", unsafe_allow_html=True)
                with pg_cols[2]:
                    if st.button("Next →", disabled=st.session_state.hist_page >= total_pages - 1, width='stretch', key="hist_next"):
                        st.session_state.hist_page += 1
                        st.rerun()

            # Action buttons per blog (below grid, using selectbox + buttons)
            st.markdown(f"<div style='height:20px'></div>", unsafe_allow_html=True)
            st.markdown(f'<div style="color:{MUTED};font-size:10.5px;font-weight:700;letter-spacing:1px;text-transform:uppercase;margin-bottom:10px">🔧 Manage Blogs</div>', unsafe_allow_html=True)

            blog_titles = [f"{i+1}. {(h['title'][:50]+'…') if len(h['title'])>50 else h['title']}" for i, h in enumerate(page_items)]
            selected_idx = st.selectbox("Select a blog", options=range(len(page_items)),
                format_func=lambda i: blog_titles[i], label_visibility="collapsed")

            selected = page_items[selected_idx]
            act1, act2, act3 = st.columns([1, 2, 1])

            with act1:
                st.markdown('<div class="view-btn">', unsafe_allow_html=True)
                if st.button("👁️ View Blog", key="view_sel", width='stretch'):
                    st.session_state.view_blog = selected
                    st.rerun()
                st.markdown('</div>', unsafe_allow_html=True)

            with act2:
                _sbase = selected["fn"].replace(".md", "")
                dl_s1, dl_s2, dl_s3 = st.columns(3)
                with dl_s1:
                    st.download_button("📄 .md", data=selected["content"],
                        file_name=selected["fn"], mime="text/markdown",
                        width='stretch', key="dl_sel_md")
                with dl_s2:
                    st.download_button("🌐 .html", data=md_to_html(selected["content"], _sbase),
                        file_name=_sbase+".html", mime="text/html",
                        width='stretch', key="dl_sel_html")
                with dl_s3:
                    st.download_button("📝 .docx", data=md_to_docx(selected["content"], _sbase),
                        file_name=_sbase+".docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        width='stretch', key="dl_sel_docx")

            with act3:
                st.markdown('<div class="del-btn">', unsafe_allow_html=True)
                if st.button("🗑️ Delete", key="del_sel", width='stretch'):
                    try:
                        os.remove(selected["path"])
                        if selected["img_path"]:
                            os.remove(selected["img_path"])
                        meta_path = selected["path"].replace("blog_", "meta_").replace(".md", ".json")
                        if os.path.exists(meta_path):
                            os.remove(meta_path)
                        load_history.clear()
                        get_thumb_b64.clear()
                        st.success("Blog deleted.")
                        st.rerun()
                    except Exception as de:
                        st.error(f"Delete failed: {de}")
                st.markdown('</div>', unsafe_allow_html=True)

    st.stop()

# ── init view_blog for generate page ──
if "view_blog" not in st.session_state:
    st.session_state.view_blog = None


# ══════════════════════════════════════════
#  PAGE: GENERATE
# ══════════════════════════════════════════

# HERO
st.markdown(f"""
<div class="hero">
  <div class="hero-badge">✦ &nbsp;AI POWERED</div>
  <div class="hero-title">
    <span class="g1">Video</span> to <span class="g2">Blog</span>
  </div>
 
  <div class="hero-tags">
    <span class="hero-tag">⚡ Instant Blog Generation</span>
    <span class="hero-tag">🌐 10 Languages</span>
    <span class="hero-tag">🎨 AI Cover Image</span>
    <span class="hero-tag">📥 Download as Markdown</span>
  </div>
</div>
""", unsafe_allow_html=True)

# FORM CARD
st.markdown('<div class="fcard">', unsafe_allow_html=True)

st.markdown(f"""
<div class="flabel">
  <svg width="16" height="12" viewBox="0 0 16 12" fill="none">
    <rect width="16" height="12" rx="2.5" fill="#FF0000"/>
    <polygon points="6.5,2.5 6.5,9.5 12,6" fill="white"/>
  </svg>
  Video URL
</div>
""", unsafe_allow_html=True)

video_url = st.text_input("url", placeholder="https://www.youtube.com/watch?v=...", label_visibility="collapsed")

# Live thumbnail preview
if video_url.strip():
    vid_preview = extract_video_id(video_url.strip())
    if vid_preview:
        thumb = get_thumbnail_url(vid_preview)
        st.markdown(f"""
        <div class="thumb-box" style="margin-top:12px">
          <img src="{thumb}" onerror="this.style.display='none';this.nextSibling.style.display='flex'"
               style="width:100%;border-radius:12px;display:block">
          <div style="display:none;background:rgba(124,58,237,.1);border-radius:12px;
               padding:20px;text-align:center;color:{MUTED};font-size:13px">
            🎬 Video preview unavailable
          </div>
        </div>
        """, unsafe_allow_html=True)

st.markdown('<div class="fdivider"></div>', unsafe_allow_html=True)
st.markdown('<div class="fsublabel">⚙️ &nbsp;Customize Output</div>', unsafe_allow_html=True)

st.markdown(f"<p style='color:{TEXT};font-weight:700;margin-bottom:4px'>🌐 Language</p>", unsafe_allow_html=True)
blog_language = st.radio("Language", 
    ["English","Urdu","Hindi","French","Spanish","Arabic","German","Portuguese","Turkish","Bengali"],
    horizontal=True, label_visibility="collapsed")

st.markdown(f"<p style='color:{TEXT};font-weight:700;margin-bottom:4px'>🎨 Tone</p>", unsafe_allow_html=True)
blog_tone = st.radio("Tone",
    ["Professional","Casual & Friendly","Educational","Storytelling","Technical","Persuasive"],
    horizontal=True, label_visibility="collapsed")

st.markdown('<div class="fdivider"></div>', unsafe_allow_html=True)
st.markdown('<div class="gen-btn">', unsafe_allow_html=True)
generate_btn = st.button("🚀  Generate Blog Post", width='stretch', key="gen_btn")
st.markdown('</div>', unsafe_allow_html=True)
st.markdown('</div>', unsafe_allow_html=True)  # /fcard


# ══════════════════════════════════════════
#  GENERATE LOGIC
# ══════════════════════════════════════════
if generate_btn:
    steps_ph = st.empty()
    if not video_url.strip():
        st.warning("⚠️ Please enter a video URL first.")
    else:
        vid = extract_video_id(video_url)
        if not vid:
            st.error("❌ Couldn't extract video ID. Please check the URL.")
        else:
            try:
                with st.spinner("📥 Fetching transcript..."):
                    steps_ph.markdown(render_steps(active="Transcript"), unsafe_allow_html=True)
                    data = fetch_transcript_with_proxy(vid)
                    full_text = " ".join([e.text for e in data])
                    word_count = len(full_text.split())

                with st.spinner("🧠 Building semantic context..."):
                    steps_ph.markdown(render_steps(active="Context", done=["Transcript"]), unsafe_allow_html=True)
                    splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=150)
                    docs = splitter.create_documents([full_text])
                    embeddings = GoogleGenerativeAIEmbeddings(
                        model="models/gemini-embedding-001", google_api_key=GEMINI_API_KEY)
                    vectorstore = FAISS.from_documents(docs, embeddings)
                    retriever = vectorstore.as_retriever(search_kwargs={"k": 6})

                with st.spinner("✍️ Writing blog with Gemini 2.5 Flash..."):
                    steps_ph.markdown(render_steps(active="Blog", done=["Transcript","Context"]), unsafe_allow_html=True)
                    llm = ChatGoogleGenerativeAI(
                        model="gemini-2.5-flash", google_api_key=GEMINI_API_KEY, temperature=0.75)
                    ctx_docs = retriever.invoke("main topics, key insights, important points")
                    context  = "\n\n".join([d.page_content for d in ctx_docs])
                    tone_desc = TONE_MAP.get(blog_tone, "professional")
                    prompt = f"""You are an expert blog writer. Write a high-quality blog post in **{blog_language}** with a **{tone_desc}** tone.

Source material:
---
{context}
---

Structure:
1. Compelling SEO-friendly title (# heading)
2. Hook introduction (2-3 paragraphs)
3. 4-5 sections with ## headings and ### sub-headings where needed
4. Key takeaways with bullet points
5. Strong conclusion with call-to-action

Requirements:
- Write entirely in {blog_language}
- Use proper Markdown formatting
- Minimum 700 words
- Do NOT mention this is based on a YouTube video
- Make it engaging and publication-ready"""
                    response = llm.invoke(prompt)
                    blog_content = response.content
                    
                    # Generate hashtags
                    title_match = re.search(r'^#\s+(.+)', blog_content, re.MULTILINE)
                    blog_title = title_match.group(1) if title_match else "Blog Post"
                    hashtags = generate_hashtags(blog_content, blog_title, llm)

                cover_image = None
                with st.spinner("🎨 Generating cover image with Imagen 4.0..."):
                    steps_ph.markdown(render_steps(active="Cover", done=["Transcript","Context","Blog"]), unsafe_allow_html=True)
                    try:
                        gc = genai.Client(api_key=GEMINI_API_KEY)
                        img_resp = gc.models.generate_images(
                            model="imagen-4.0-generate-001",
                            prompt=f"Professional modern blog cover for: '{blog_title}'. Cinematic lighting, clean composition, no text.",
                            config=types.GenerateImagesConfig(
                                number_of_images=1,
                                aspect_ratio="16:9"
                            )
                        )
                        cover_image = Image.open(io.BytesIO(img_resp.generated_images[0].image.image_bytes))
                    except Exception as ie:
                        st.warning(f"⚠️ Cover image skipped: {ie}")

                ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                blog_fn = f"blog_{ts}.md"
                with open(os.path.join(SAVE_DIR, blog_fn), "w", encoding="utf-8") as f:
                    f.write(blog_content)
                save_meta(ts, blog_language, blog_tone, hashtags)
                load_history.clear()  # invalidate cache so history shows new blog

                cover_bytes, img_fn = None, None
                if cover_image:
                    img_fn = f"cover_{ts}.png"
                    cover_image.save(os.path.join(SAVE_DIR, img_fn), format="PNG")
                    buf = io.BytesIO()
                    cover_image.save(buf, format="PNG")
                    cover_bytes = buf.getvalue()

                st.session_state.blog_content  = blog_content
                st.session_state.cover_bytes   = cover_bytes
                st.session_state.blog_filename = blog_fn
                st.session_state.img_filename  = img_fn
                st.session_state.word_count    = word_count
                st.session_state.blog_language = blog_language
                st.session_state.blog_tone     = blog_tone
                st.session_state.hashtags      = hashtags

                steps_ph.markdown(render_steps(done=["Transcript","Context","Blog","Cover","Done"]), unsafe_allow_html=True)

            except TranscriptsDisabled:
                steps_ph.markdown(render_steps(), unsafe_allow_html=True)
                st.error("🚫 Transcripts are disabled for this video.")
            except NoTranscriptFound:
                steps_ph.markdown(render_steps(), unsafe_allow_html=True)
                st.error("❌ No transcript found for this video.")
            except Exception as e:
                steps_ph.markdown(render_steps(), unsafe_allow_html=True)
                st.error(f"⚠️ Something went wrong: {e}")


# ══════════════════════════════════════════
#  RESULT SECTION
# ══════════════════════════════════════════
if st.session_state.blog_content:
    bc  = st.session_state.blog_content
    cb  = st.session_state.cover_bytes
    bf  = st.session_state.blog_filename
    imf = st.session_state.img_filename
    wc  = st.session_state.word_count
    bl  = st.session_state.blog_language
    bt  = st.session_state.blog_tone
    bw  = len(bc.split())
    rt  = max(1, bw // 200)

    # Success banner
    st.markdown(f"""
    <div class="sbanner">
      <div class="sbanner-icon">✅</div>
      <div>
        <div class="sbanner-title">Blog generated successfully!</div>
        <div class="sbanner-sub">Auto-saved → outputs/{bf}{f"  &  outputs/{imf}" if imf else ""}</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # Stats chips
    st.markdown(f"""
    <div class="stats">
      <div class="schip">📄 Transcript <strong>{wc:,} words</strong></div>
      <div class="schip">✍️ Blog <strong>{bw:,} words</strong></div>
      <div class="schip">⏱️ <strong>~{rt} min read</strong></div>
      <div class="schip">🌐 <strong>{bl}</strong></div>
      <div class="schip">🎨 <strong>{bt}</strong></div>
    </div>
    """, unsafe_allow_html=True)

    # Hashtags section
    if st.session_state.hashtags:
        hashtags_html = '<div style="margin-top:16px;padding:14px 16px;background:' + CARD2 + ';border:1px solid ' + BORDER_S + ';border-radius:12px">'
        hashtags_html += '<div style="color:' + MUTED + ';font-size:10.5px;font-weight:700;letter-spacing:1px;text-transform:uppercase;margin-bottom:10px">🏷️ Suggested Hashtags</div>'
        hashtags_html += '<div style="display:flex;flex-wrap:wrap;gap:8px">'
        for tag in st.session_state.hashtags:
            hashtags_html += '<span style="background:' + CHIP + ';border:1px solid ' + CHIP_B + ';color:' + TEXT2 + ';padding:6px 12px;border-radius:100px;font-size:12px;font-weight:500;cursor:pointer" onclick="navigator.clipboard.writeText(\'' + tag + '\')">' + tag + '</span>'
        hashtags_html += '</div></div>'
        st.markdown(hashtags_html, unsafe_allow_html=True)

    # Result card
    st.markdown('<div class="rcard">', unsafe_allow_html=True)
    st.markdown(f"""
    <div class="rhead">
      <div class="ricon">📝</div>
      <div>
        <div class="rtitle">Your Generated Blog Post</div>
        <div class="rsub">Ready to publish — preview, copy, or download below</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # Cover image
    if cb:
        st.image(io.BytesIO(cb), width='stretch')
        st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)

    # Tabs: Preview | Raw Markdown
    tab_prev, tab_raw = st.tabs(["📖 Preview", "📝 Raw Markdown"])

    with tab_prev:
        # Copy button — store content in a JS variable to avoid HTML injection issues
        safe = bc.replace("\\", "\\\\").replace("`", "\\`").replace("$", "\\$")
        copy_btn_html = (
            '<script>window._blogContent = `' + safe + '`;</script>'
            '<div style="display:flex;justify-content:flex-end;margin-bottom:14px">'
            '<button class="cpbtn" onclick="navigator.clipboard.writeText(window._blogContent)'
            '.then(()=>{this.innerHTML=\'✅ Copied!\';setTimeout(()=>this.innerHTML=\'📋 Copy Blog\',2200);})">'
            '📋 Copy Blog</button></div>'
        )
        st.markdown(copy_btn_html, unsafe_allow_html=True)
        st.markdown('<div class="blog">', unsafe_allow_html=True)
        st.markdown(bc)
        st.markdown('</div>', unsafe_allow_html=True)

    with tab_raw:
        st.code(bc, language="markdown")

    # Downloads
    st.markdown(f"""
    <div style="height:1px;background:{BORDER_S};margin:22px 0 16px"></div>
    <div style="color:{MUTED};font-size:10.5px;font-weight:700;letter-spacing:1px;
    text-transform:uppercase;margin-bottom:12px">⬇️ &nbsp;Download</div>
    """, unsafe_allow_html=True)

    _base = bf.replace(".md", "")
    if cb:
        d1, d2, d3, d4 = st.columns(4)
        with d1:
            st.download_button("📄 .md", data=bc, file_name=bf,
                mime="text/markdown", width='stretch', key="dl_blog")
        with d2:
            st.download_button("🌐 .html", data=md_to_html(bc, _base),
                file_name=_base+".html", mime="text/html", width='stretch', key="dl_html")
        with d3:
            st.download_button("📝 .docx", data=md_to_docx(bc, _base),
                file_name=_base+".docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                width='stretch', key="dl_docx")
        with d4:
            st.download_button("🖼️ .png", data=cb, file_name=imf,
                mime="image/png", width='stretch', key="dl_img")
    else:
        d1, d2, d3 = st.columns(3)
        with d1:
            st.download_button("📄 .md", data=bc, file_name=bf,
                mime="text/markdown", width='stretch', key="dl_blog")
        with d2:
            st.download_button("🌐 .html", data=md_to_html(bc, _base),
                file_name=_base+".html", mime="text/html", width='stretch', key="dl_html")
        with d3:
            st.download_button("📝 .docx", data=md_to_docx(bc, _base),
                file_name=_base+".docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                width='stretch', key="dl_docx")

    st.markdown('</div>', unsafe_allow_html=True)  # /rcard

    # Share buttons
    _blog_title = st.session_state.get("video_title", "Check out this blog")
    _share_text = f"{_blog_title} — generated with VidBlog AI"
    _li_url  = f"https://www.linkedin.com/sharing/share-offsite/?url=https://vidblog.ai"
    _tw_url  = f"https://twitter.com/intent/tweet?text={requests.utils.quote(_share_text)}"
    _fb_url  = f"https://www.facebook.com/sharer/sharer.php?u=https://vidblog.ai&quote={requests.utils.quote(_share_text)}"
    _wa_url  = f"https://api.whatsapp.com/send?text={requests.utils.quote(_share_text)}"

    if "show_share" not in st.session_state:
        st.session_state.show_share = False

    st.markdown("<div style='margin-top:16px'>", unsafe_allow_html=True)
    _scol, _ = st.columns([1, 3])
    with _scol:
        if st.button("🔗 Share", key="share_toggle", width='stretch'):
            st.session_state.show_share = not st.session_state.show_share

    if st.session_state.show_share:
        st.markdown(f"""
        <div style="display:flex;gap:10px;flex-wrap:wrap;margin-top:12px">
          <a href="{_li_url}" target="_blank" style="display:inline-flex;align-items:center;gap:7px;padding:9px 18px;border-radius:10px;background:#0a66c2;color:#fff;font-size:13px;font-weight:600;text-decoration:none;font-family:'Inter',sans-serif;">
            <svg width="16" height="16" viewBox="0 0 24 24" fill="white"><path d="M20.447 20.452h-3.554v-5.569c0-1.328-.027-3.037-1.852-3.037-1.853 0-2.136 1.445-2.136 2.939v5.667H9.351V9h3.414v1.561h.046c.477-.9 1.637-1.85 3.37-1.85 3.601 0 4.267 2.37 4.267 5.455v6.286zM5.337 7.433a2.062 2.062 0 01-2.063-2.065 2.064 2.064 0 112.063 2.065zm1.782 13.019H3.555V9h3.564v11.452zM22.225 0H1.771C.792 0 0 .774 0 1.729v20.542C0 23.227.792 24 1.771 24h20.451C23.2 24 24 23.227 24 22.271V1.729C24 .774 23.2 0 22.222 0h.003z"/></svg>
            LinkedIn
          </a>
          <a href="{_tw_url}" target="_blank" style="display:inline-flex;align-items:center;gap:7px;padding:9px 18px;border-radius:10px;background:#000;color:#fff;font-size:13px;font-weight:600;text-decoration:none;font-family:'Inter',sans-serif;">
            <svg width="16" height="16" viewBox="0 0 24 24" fill="white"><path d="M18.244 2.25h3.308l-7.227 8.26 8.502 11.24H16.17l-4.714-6.231-5.401 6.231H2.744l7.737-8.835L1.254 2.25H8.08l4.253 5.622zm-1.161 17.52h1.833L7.084 4.126H5.117z"/></svg>
            X (Twitter)
          </a>
          <a href="{_fb_url}" target="_blank" style="display:inline-flex;align-items:center;gap:7px;padding:9px 18px;border-radius:10px;background:#1877f2;color:#fff;font-size:13px;font-weight:600;text-decoration:none;font-family:'Inter',sans-serif;">
            <svg width="16" height="16" viewBox="0 0 24 24" fill="white"><path d="M24 12.073c0-6.627-5.373-12-12-12s-12 5.373-12 12c0 5.99 4.388 10.954 10.125 11.854v-8.385H7.078v-3.47h3.047V9.43c0-3.007 1.792-4.669 4.533-4.669 1.312 0 2.686.235 2.686.235v2.953H15.83c-1.491 0-1.956.925-1.956 1.874v2.25h3.328l-.532 3.47h-2.796v8.385C19.612 23.027 24 18.062 24 12.073z"/></svg>
            Facebook
          </a>
          <a href="{_wa_url}" target="_blank" style="display:inline-flex;align-items:center;gap:7px;padding:9px 18px;border-radius:10px;background:#25d366;color:#fff;font-size:13px;font-weight:600;text-decoration:none;font-family:'Inter',sans-serif;">
            <svg width="16" height="16" viewBox="0 0 24 24" fill="white"><path d="M17.472 14.382c-.297-.149-1.758-.867-2.03-.967-.273-.099-.471-.148-.67.15-.197.297-.767.966-.94 1.164-.173.199-.347.223-.644.075-.297-.15-1.255-.463-2.39-1.475-.883-.788-1.48-1.761-1.653-2.059-.173-.297-.018-.458.13-.606.134-.133.298-.347.446-.52.149-.174.198-.298.298-.497.099-.198.05-.371-.025-.52-.075-.149-.669-1.612-.916-2.207-.242-.579-.487-.5-.669-.51-.173-.008-.371-.01-.57-.01-.198 0-.52.074-.792.372-.272.297-1.04 1.016-1.04 2.479 0 1.462 1.065 2.875 1.213 3.074.149.198 2.096 3.2 5.077 4.487.709.306 1.262.489 1.694.625.712.227 1.36.195 1.871.118.571-.085 1.758-.719 2.006-1.413.248-.694.248-1.289.173-1.413-.074-.124-.272-.198-.57-.347m-5.421 7.403h-.004a9.87 9.87 0 01-5.031-1.378l-.361-.214-3.741.982.998-3.648-.235-.374a9.86 9.86 0 01-1.51-5.26c.001-5.45 4.436-9.884 9.888-9.884 2.64 0 5.122 1.03 6.988 2.898a9.825 9.825 0 012.893 6.994c-.003 5.45-4.437 9.884-9.885 9.884m8.413-18.297A11.815 11.815 0 0012.05 0C5.495 0 .16 5.335.157 11.892c0 2.096.547 4.142 1.588 5.945L.057 24l6.305-1.654a11.882 11.882 0 005.683 1.448h.005c6.554 0 11.89-5.335 11.893-11.893a11.821 11.821 0 00-3.48-8.413z"/></svg>
            WhatsApp
          </a>
        </div>
        """, unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)

    # Footer CTA
    st.markdown(f"""
    <div class="cta-footer">
      <p>Want to generate another blog?</p>
      <span>Paste a new video URL above and click Generate again — or check 📂 History to revisit past blogs.</span>
    </div>
    """, unsafe_allow_html=True)
